#!/usr/bin/env python3
"""Render the markdown documentation set as branded Word documents.

Produces one .docx per source file plus a single combined handbook with a contents
page, into ~/Documents/SAP/Supply_Chain_Ontology_Docs/.

Why a bespoke renderer rather than pandoc: the output has to match the two existing
Word deliverables (management summary, demo scripts) so the set reads as one family
of documents. That means docx_kit's palette, page setup, navy table headers and
`keep_with_next` behaviour — none of which survive a generic converter.

Markdown handled: ATX headings, fenced code, bullets with one nesting level,
numbered lists, blockquotes, horizontal rules, GFM tables, and inline **bold**,
`code` and [links](url).

The table parser accepts rows with or without leading and trailing pipes. GFM does
not require them and 139 rows across this doc set omit them, so a parser that
insists on them silently drops content.
"""

import pathlib
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

sys.path.insert(0, str(pathlib.Path(__file__).parent))
from docx_kit import (  # noqa: E402
    BLUE_HEX, GREY, LIGHT_HEX, NAVY_HEX, SAP_NAVY, SNOW_BLUE,
    fixed, h1, h2, no_split, setup_page, shade,
)

REPO = pathlib.Path(__file__).resolve().parent.parent
OUT = pathlib.Path.home() / "Documents" / "SAP" / "Supply_Chain_Ontology_Docs"

# Order matters: this is the reading order of the handbook.
SOURCES = [
    ("README.md",                   "Overview",              "00"),
    ("docs/01-concepts.md",         "Concepts",              "01"),
    ("docs/02-architecture.md",     "Architecture",          "02"),
    ("docs/03-build-steps.md",      "Build steps",           "03"),
    ("docs/04-execution.md",        "Execution",             "04"),
    ("docs/05-findings.md",         "Findings",              "05"),
    ("docs/06-references.md",       "References",            "06"),
    ("docs/07-scenario-modelling.md", "Scenario modelling",  "07"),
]

MONO = "Consolas"
CODE_FILL = "F4F6F8"


# --------------------------------------------------------------------- inline

def rich_md(paragraph, text, size, base_color=None):
    """Write inline markdown into a paragraph.

    Extends docx_kit.rich() with `code` spans and [links](url). Kept here rather
    than replacing rich() because the management summary and demo scripts are
    already verified against its exact behaviour.
    """
    # Links become "text (url)" only when the url adds information — a bare
    # self-linking url would read twice.
    def unlink(m):
        label, url = m.group(1), m.group(2)
        return label if label.rstrip("/") == url.rstrip("/") else f"{label} ({url})"

    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", unlink, str(text))

    def emit(chunk, bold):
        """Emit one chunk, splitting out `code` spans and *italic* runs.

        Code is split before italic so a path like `a*b` stays literal.
        """
        for part in re.split(r"(`[^`]+`)", chunk):
            if not part:
                continue
            if part.startswith("`") and part.endswith("`"):
                r = paragraph.add_run(part[1:-1])
                r.font.name = MONO
                r.font.size = Pt(size - 0.5)
                r.font.bold = bold
                r.font.color.rgb = RGBColor(0x1B, 0x5E, 0x8C)
                continue
            for piece in re.split(r"(?<!\*)(\*[^*\n]+?\*)(?!\*)", part):
                if not piece:
                    continue
                ital = piece.startswith("*") and piece.endswith("*") and len(piece) > 2
                r = paragraph.add_run(piece[1:-1] if ital else piece)
                r.font.size = Pt(size)
                r.font.bold = bold
                r.font.italic = ital or None
                if base_color is not None:
                    r.font.color.rgb = base_color

    # Bold first, then code inside each side of it. A code span can legitimately
    # contain ** (glob patterns like client/**), so the code split has to happen
    # within the bold segments rather than competing with them at the same level.
    for seg in re.split(r"(\*\*[^*]+?\*\*)", text):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            emit(seg[2:-2], True)
        else:
            emit(seg, False)
    return paragraph


def heading(doc, text, size, color, before, after):
    """A heading that honours inline markdown, bolded throughout."""
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    rich_md(p, text, size, color)
    for r in p.runs:
        r.font.bold = True
        r.font.color.rgb = color
    return p


def para(doc, text, size=10, after=6, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    rich_md(p, text, size, color)
    for r in p.runs:
        r.italic = italic
    return p


def bullet_md(doc, text, size=10, level=0):
    p = doc.add_paragraph(style="List Bullet 2" if level else "List Bullet")
    p.paragraph_format.space_after = Pt(2)
    rich_md(p, text, size)
    return p


def numbered_md(doc, text, size=10):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    rich_md(p, text, size)
    return p


def code_block(doc, lines):
    """A fenced block as a single shaded, monospaced, unsplittable cell."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    fixed(t, [6.9])
    cell = t.rows[0].cells[0]
    shade(cell, CODE_FILL)
    no_split(t.rows[0])
    cell.paragraphs[0].text = ""
    for i, ln in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(ln if ln.strip() else " ")
        r.font.name = MONO
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x22, 0x33, 0x44)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def quote(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    rich_md(p, " ".join(lines), 9.5, GREY)
    for r in p.runs:
        r.italic = True
    return p


def rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("─" * 58)
    r.font.size = Pt(6)
    r.font.color.rgb = RGBColor(0xC8, 0xD4, 0xDE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def md_table(doc, rows):
    """Render a parsed GFM table. rows[0] is the header."""
    headers, data = rows[0], rows[1:]
    ncol = max(len(r) for r in rows)
    headers = headers + [""] * (ncol - len(headers))
    data = [r + [""] * (ncol - len(r)) for r in data]

    # First column tends to be a label and stays narrow; the rest share the space.
    total = 6.9
    first = min(2.0, max(1.1, total / ncol))
    widths = [first] + [(total - first) / (ncol - 1)] * (ncol - 1) if ncol > 1 else [total]

    t = doc.add_table(rows=1 + len(data), cols=ncol)
    t.style = "Table Grid"
    fixed(t, widths)

    for j, htext in enumerate(headers):
        c = t.rows[0].cells[j]
        shade(c, NAVY_HEX)
        c.paragraphs[0].text = ""
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(re.sub(r"[*`]", "", htext))
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    no_split(t.rows[0], header=True)

    for i, row in enumerate(data):
        for j, val in enumerate(row):
            c = t.rows[i + 1].cells[j]
            if i % 2:
                shade(c, LIGHT_HEX)
            c.paragraphs[0].text = ""
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            rich_md(p, val, 8.5)
        no_split(t.rows[i + 1])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ---------------------------------------------------------------------- parse

# A cell may contain an escaped pipe (\|), which must survive the row split —
# docs/06-references.md has `BASE_PATH \|\| '/'` inside a code span, and splitting
# on it tears the span in half across two cells.
PIPE_SENTINEL = "\x00PIPE\x00"


def split_row(line):
    """Split a GFM table row, tolerating missing pipes and honouring \| escapes."""
    s = line.strip().replace("\\|", PIPE_SENTINEL)
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip().replace(PIPE_SENTINEL, "|") for c in s.split("|")]


def is_separator(line):
    """`---|:--:|---` style row that marks the end of a table header."""
    cells = split_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c != "")


def render(doc, text, base_level=0):
    """Walk the markdown and emit Word content. base_level shifts heading depth."""
    lines = text.split("\n")
    i, n = 0, len(lines)
    buf: list[str] = []

    def flush():
        nonlocal buf
        if buf:
            para(doc, " ".join(buf))
            buf = []

    while i < n:
        ln = lines[i]

        # fenced code
        if ln.lstrip().startswith("```"):
            flush()
            i += 1
            block = []
            while i < n and not lines[i].lstrip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            if block:
                code_block(doc, block)
            continue

        # table: a pipe row followed by a separator row
        if ln.count("|") >= 1 and i + 1 < n and is_separator(lines[i + 1]):
            flush()
            rows = [split_row(ln)]
            i += 2
            while i < n and lines[i].count("|") >= 1 and lines[i].strip():
                rows.append(split_row(lines[i]))
                i += 1
            md_table(doc, rows)
            continue

        # headings
        m = re.match(r"^(#{1,4})\s+(.*)$", ln)
        if m:
            flush()
            lvl = len(m.group(1)) + base_level
            title = m.group(2).strip()
            # Headings can contain [links](url) and `code` — several H3s in the
            # references are a linked repo name — so they render through rich_md
            # rather than as one plain run.
            if lvl <= 1:
                heading(doc, title, 15, SAP_NAVY, 18, 6)
            elif lvl == 2:
                heading(doc, title, 11.5, SAP_NAVY, 12, 3)
            else:
                heading(doc, title, 10, SNOW_BLUE, 8, 2)
            i += 1
            continue

        # horizontal rule (a bare --- that is not a table separator)
        if re.fullmatch(r"\s*([-*_])\1{2,}\s*", ln):
            flush()
            rule(doc)
            i += 1
            continue

        # blockquote
        if ln.lstrip().startswith(">"):
            flush()
            q = []
            while i < n and lines[i].lstrip().startswith(">"):
                q.append(lines[i].lstrip()[1:].strip())
                i += 1
            quote(doc, [x for x in q if x])
            continue

        # bullets, one nesting level
        m = re.match(r"^(\s*)[-*]\s+(.*)$", ln)
        if m:
            flush()
            indent, txt = m.group(1), m.group(2)
            # continuation lines of the same bullet
            i += 1
            while i < n and lines[i].strip() and not re.match(
                    r"^(\s*)([-*]|\d+\.)\s+", lines[i]) and not lines[i].startswith("#") \
                    and "|" not in lines[i] and not lines[i].lstrip().startswith("```"):
                txt += " " + lines[i].strip()
                i += 1
            bullet_md(doc, txt, level=1 if len(indent) >= 2 else 0)
            continue

        # numbered
        m = re.match(r"^\s*\d+\.\s+(.*)$", ln)
        if m:
            flush()
            txt = m.group(1)
            i += 1
            while i < n and lines[i].strip() and not re.match(
                    r"^(\s*)([-*]|\d+\.)\s+", lines[i]) and not lines[i].startswith("#") \
                    and "|" not in lines[i] and not lines[i].lstrip().startswith("```"):
                txt += " " + lines[i].strip()
                i += 1
            numbered_md(doc, txt)
            continue

        if not ln.strip():
            flush()
            i += 1
            continue

        buf.append(ln.strip())
        i += 1

    flush()


# --------------------------------------------------------------------- output

def cover(doc, title, subtitle, part=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Supply Chain Ontology")
    r.font.size = Pt(13)
    r.font.color.rgb = SNOW_BLUE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = SAP_NAVY

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle + (f"  ·  part {part}" if part else ""))
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY


def build_one(src, title, num):
    text = (REPO / src).read_text()
    # the first H1 becomes the cover title, so it is not repeated in the body
    body_text = re.sub(r"^#\s+.*\n", "", text, count=1)
    doc = Document()
    setup_page(doc)
    cover(doc, title, f"Generated from {src}", num)
    render(doc, body_text)
    out = OUT / f"{num}_{title.replace(' ', '_')}.docx"
    doc.save(out)
    return out, len(body_text.split("\n"))


def build_handbook():
    doc = Document()
    setup_page(doc)
    cover(doc, "Documentation handbook",
          "Every document in one volume, in reading order")

    h1(doc, "Contents")
    t = doc.add_table(rows=len(SOURCES), cols=2)
    t.style = "Table Grid"
    fixed(t, [1.1, 5.8])
    for i, (src, title, num) in enumerate(SOURCES):
        c0, c1 = t.rows[i].cells
        shade(c0, LIGHT_HEX)
        for cell, val, bold in ((c0, num, True), (c1, title, False)):
            cell.paragraphs[0].text = ""
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(9)
            r.font.bold = bold
            r.font.color.rgb = SAP_NAVY if bold else None
        no_split(t.rows[i])
    doc.add_paragraph()

    for src, title, num in SOURCES:
        doc.add_page_break()
        h1(doc, f"{num} · {title}", size=17)
        text = re.sub(r"^#\s+.*\n", "", (REPO / src).read_text(), count=1)
        # shift headings down one level so the part title stays dominant
        render(doc, text, base_level=1)

    out = OUT / "Supply_Chain_Ontology_Documentation_Handbook.docx"
    doc.save(out)
    return out


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    total = 0
    for src, title, num in SOURCES:
        if not (REPO / src).exists():
            sys.exit(f"missing source: {src}")
        out, nlines = build_one(src, title, num)
        kb = out.stat().st_size // 1024
        total += out.stat().st_size
        print(f"  {out.name:52} {kb:>4} KB   from {nlines} md lines")

    hb = build_handbook()
    total += hb.stat().st_size
    print(f"  {hb.name:52} {hb.stat().st_size // 1024:>4} KB   combined")
    print(f"\n  {len(SOURCES) + 1} documents, {total // 1024} KB total")
    print(f"  in {OUT}")


if __name__ == "__main__":
    main()
