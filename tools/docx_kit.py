#!/usr/bin/env python3
"""Shared python-docx helpers for the project's Word deliverables.

Extracted so the management summary and the demo scripts cannot drift apart in
layout or branding. Both import from here.

The two non-obvious pieces are documented at their definitions: fixed table
layout has to be expressed through w:tblGrid, and inline bold has to be parsed
per segment rather than per cell.
"""
import re

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SNOW_BLUE = RGBColor(0x29, 0xB5, 0xE8)
SAP_NAVY = RGBColor(0x1B, 0x3A, 0x57)
GREY = RGBColor(0x5A, 0x6A, 0x7A)
RED = RGBColor(0xC0, 0x28, 0x28)
GREEN = RGBColor(0x1B, 0x7F, 0x4B)
AMBER = RGBColor(0xB4, 0x6A, 0x00)

NAVY_HEX = "1B3A57"
BLUE_HEX = "29B5E8"
LIGHT_HEX = "EEF4F8"


def fixed(t, widths):
    """Pin real column widths.

    In fixed layout Word takes widths from w:tblGrid/w:gridCol and ignores w:tcW,
    so setting cell width alone silently yields equal columns.
    """
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    t._tbl.tblPr.append(lay)
    grid = t._tbl.find(qn("w:tblGrid"))
    for col, wid in zip(grid.findall(qn("w:gridCol")), widths):
        col.set(qn("w:w"), str(int(wid * 1440)))
    for row in t.rows:
        for cell, wid in zip(row.cells, widths):
            cell.width = Inches(wid)


def no_split(row, header=False):
    """Keep a row whole across a page break; repeat header rows on each page."""
    pr = row._tr.get_or_add_trPr()
    pr.append(OxmlElement("w:cantSplit"))
    if header:
        pr.append(OxmlElement("w:tblHeader"))


def shade(cell, hexcolor):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(sh)


def rich(paragraph, text, size, base_color=None):
    """Write text into a paragraph, honouring **bold** segments.

    Parsed per segment rather than per cell: "**Real** — from SAP data" needs the
    lead bolded and the remainder plain, and a whole-string test renders the
    asterisks literally.
    """
    for seg in re.split(r"(\*\*[^*]+\*\*)", str(text)):
        if not seg:
            continue
        r = paragraph.add_run(seg.strip("*") if seg.startswith("**") else seg)
        r.font.size = Pt(size)
        r.font.bold = seg.startswith("**")
        if base_color is not None:
            r.font.color.rgb = base_color
    return paragraph


def h1(doc, text, size=15, color=SAP_NAVY, before=18, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = color
    return p


def h2(doc, text, size=11.5, color=SAP_NAVY):
    return h1(doc, text, size=size, color=color, before=12, after=3)


def body(doc, text, size=10, italic=False, color=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    rich(p, text, size, color)
    for r in p.runs:
        r.italic = italic
    return p


def bullet(doc, text, size=10, style="List Bullet"):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    rich(p, text, size)
    return p


def table(doc, headers, rows, widths, size=9, header_fill=NAVY_HEX,
          align_right=(), zebra=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0]
    for i, htext in enumerate(headers):
        c = hdr.cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(htext)
        r.font.size = Pt(size)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c, header_fill)
    no_split(hdr, header=True)

    for n, row in enumerate(rows):
        cells = t.add_row()
        no_split(cells)
        for i, val in enumerate(row):
            c = cells.cells[i]
            c.text = ""
            p = c.paragraphs[0]
            if i in align_right:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rich(p, val, size)
            if zebra and n % 2 == 1:
                shade(c, LIGHT_HEX)
    fixed(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def callout(doc, label, text, fill=LIGHT_HEX, size=9.5):
    """A single-cell shaded box, for the line the presenter should actually say."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    c.text = ""
    p = c.paragraphs[0]
    if label:
        r = p.add_run(label + "  ")
        r.font.size = Pt(size)
        r.font.bold = True
        r.font.color.rgb = SAP_NAVY
    rich(p, text, size)
    shade(c, fill)
    fixed(t, [6.9])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def money(n):
    n = float(n)
    if abs(n) >= 1e6:
        return f"${n / 1e6:.2f}M"
    if abs(n) >= 1e3:
        return f"${n / 1e3:.0f}K"
    return f"${n:.0f}"


def setup_page(doc, margins=(0.7, 0.7, 0.8, 0.8)):
    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    s.top_margin, s.bottom_margin = Inches(margins[0]), Inches(margins[1])
    s.left_margin, s.right_margin = Inches(margins[2]), Inches(margins[3])
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10)
    return s
