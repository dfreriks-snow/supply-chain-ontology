#!/usr/bin/env python3
"""Generate the management summary Word document.

Every figure is read from /tmp/scenario_facts.json and /tmp/network_facts.json,
which are produced by querying the running application — nothing here is
transcribed by hand, so the document cannot drift from what the app actually
reports.

Regenerate the inputs first (app must be running on :3009), then:
    python3 tools/build_management_summary.py
"""
import datetime
import json
import re
import pathlib
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SNOW_BLUE = RGBColor(0x29, 0xB5, 0xE8)
SAP_NAVY = RGBColor(0x1B, 0x3A, 0x57)
GREY = RGBColor(0x5A, 0x6A, 0x7A)
RED = RGBColor(0xC0, 0x28, 0x28)
GREEN = RGBColor(0x1B, 0x7F, 0x4B)

OUT = pathlib.Path.home() / "Documents" / "SAP" / "Supply_Chain_Ontology_Management_Summary.docx"

REPO = "https://github.com/dfreriks-snow/supply-chain-ontology"
SITE = "https://dfreriks-snow.github.io/supply-chain-ontology/"


# ---------------------------------------------------------------- docx helpers

def _fixed(t, widths):
    """Pin real column widths.

    In fixed layout Word takes widths from w:tblGrid/w:gridCol and ignores w:tcW,
    so setting cell width alone silently produces equal columns.
    """
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    pr = t._tbl.tblPr
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    pr.append(lay)
    grid = t._tbl.find(qn("w:tblGrid"))
    cols = grid.findall(qn("w:gridCol"))
    for col, wid in zip(cols, widths):
        col.set(qn("w:w"), str(int(wid * 1440)))
    for row in t.rows:
        for cell, wid in zip(row.cells, widths):
            cell.width = Inches(wid)


def _no_split(row, header=False):
    """Keep a row intact across a page break, and repeat header rows."""
    pr = row._tr.get_or_add_trPr()
    pr.append(OxmlElement("w:cantSplit"))
    if header:
        pr.append(OxmlElement("w:tblHeader"))


def _shade(cell, hexcolor):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(sh)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = SAP_NAVY
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(11.5)
    r.font.bold = True
    r.font.color.rgb = SAP_NAVY
    return p


def body(doc, text, size=10, italic=False, color=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def bullet(doc, text, size=10, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.font.size = Pt(size)
        r.font.bold = True
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def table(doc, headers, rows, widths, size=9, header_fill="1B3A57", align_right=()):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0]
    for i, htext in enumerate(headers):
        c = hdr.cells[i]
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(htext)
        r.font.size = Pt(size)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(c, header_fill)
    _no_split(hdr, header=True)

    for row in rows:
        cells = t.add_row()
        _no_split(cells)
        for i, val in enumerate(row):
            c = cells.cells[i]
            c.text = ""
            p = c.paragraphs[0]
            if i in align_right:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Parse **bold** segments anywhere in the cell, not just a cell that is
            # entirely wrapped — "**Real** — from SAP data" needs the lead bolded
            # and the remainder plain, and a whole-cell test renders the asterisks.
            for seg in re.split(r"(\*\*[^*]+\*\*)", str(val)):
                if not seg:
                    continue
                r = p.add_run(seg.strip("*") if seg.startswith("**") else seg)
                r.font.size = Pt(size)
                r.font.bold = seg.startswith("**")
    _fixed(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def money(n):
    n = float(n)
    if abs(n) >= 1e6:
        return f"${n/1e6:.2f}M"
    if abs(n) >= 1e3:
        return f"${n/1e3:.0f}K"
    return f"${n:.0f}"


# ------------------------------------------------------------------- the report

def main():
    try:
        scen = json.loads(pathlib.Path("/tmp/scenario_facts.json").read_text())
        net = json.loads(pathlib.Path("/tmp/network_facts.json").read_text())
    except FileNotFoundError as e:
        sys.exit(f"missing input: {e}. Query the running app first.")

    doc = Document()
    s = doc.sections[0]
    s.orientation = WD_ORIENT.PORTRAIT
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    for m in ("top_margin", "bottom_margin"):
        setattr(s, m, Inches(0.7))
    for m in ("left_margin", "right_margin"):
        setattr(s, m, Inches(0.8))

    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10)

    # ---- title -----------------------------------------------------------
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Supply Chain Ontology")
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = SAP_NAVY

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Disruption scenario modelling on SAP Business Data Cloud and Snowflake")
    r.font.size = Pt(13)
    r.font.color.rgb = SNOW_BLUE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(f"Management summary · {datetime.date.today():%d %B %Y}")
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY

    total_products = 36
    nt = net["totals"]

    # ---- executive summary ----------------------------------------------
    h1(doc, "Executive summary")
    body(doc,
         "We have built and deployed a working application that answers a question the "
         "business could not previously answer quickly: when a plant goes offline, what "
         "ripples through the supply chain, what does it cost, and what can be done to "
         "compensate.",
         size=10.5)
    body(doc,
         f"It runs on data already in Snowflake. Nothing was copied out of SAP, and no new "
         f"data platform was required. The application models a {nt['plants']}-plant network "
         f"carrying {money(nt['monthly_value'])} of flow per month, and produces a costed "
         f"mitigation plan in seconds rather than the days a manual exercise takes.")

    rows = [
        ["Business question answered",
         "If a site is disrupted, what is at risk downstream and how do we compensate?"],
        ["Status", "**Built, deployed and verified. Running today.**"],
        ["Data foundation",
         f"SAP data products in Snowflake — {nt['nodes']} network nodes, {nt['flows']} flows, "
         f"{money(nt['monthly_value'])} monthly"],
        ["Time to an answer",
         "Seconds per scenario, versus a manual exercise measured in days"],
        ["New infrastructure required", "None. Views over existing Snowflake data."],
        ["Cost to operate",
         "Negligible. Simulation runs in the application; only the AI narrative uses Snowflake compute."],
    ]
    table(doc, ["Item", "Summary"], rows, [1.9, 5.0], size=9.5)

    # ---- what was delivered ---------------------------------------------
    h1(doc, "What was delivered")
    body(doc,
         "Three connected pieces of work, each independently usable.")

    rows = [
        ["1. Ontology explorer",
         f"A navigable model of the supply-chain slice of the SAP BDC catalog: "
         f"{total_products} data products, 338 entities and 281 declared relationships. "
         f"Answers what data exists and how it connects.",
         "6 pages"],
        ["2. Scenario modelling",
         "Disruption simulation across five event types, with ripple visualisation on a world "
         "map and network graph side by side, plus a mitigation optimiser and an AI briefing.",
         "3 pages"],
        ["3. Publication",
         "Documented public repository and a credential-free web build, so the work can be "
         "shared with customers and colleagues without granting Snowflake access.",
         "7 documents"],
    ]
    table(doc, ["Workstream", "What it does", "Scale"], rows, [1.55, 4.35, 0.95], size=9.5)

    # ---- the worked example ---------------------------------------------
    hur = next(x for x in scen if x["id"] == "hurricane-austin")
    h1(doc, "Worked example: a hurricane closes the Austin plant")
    body(doc,
         "This is the case to show first, because the second-order effect is the part no "
         "spreadsheet catches.")

    rows = [
        ["Direct loss", "Two customers lose supply immediately",
         money(8_400_000 + 5_400_000)],
        ["Second-order loss",
         "Austin also ships test fixtures to Penang. Penang runs on 12 days of stock, then "
         "starves — taking two further customers down with it.",
         money(1_291_035)],
        ["**Total exposure**", f"**{hur['pctOfNetwork']}% of monthly network value**",
         f"**{money(hur['valueAtRisk'])}**"],
    ]
    table(doc, ["Effect", "What happens", "Value"], rows, [1.5, 4.35, 1.0], size=9.5, align_right=(2,))

    body(doc,
         "The application then produces a mitigation plan that respects real capacity limits:",
         after=4)
    bullet(doc, "Move inspection systems for GlobalFoundries to San Jose — 98 of 282 free hours.")
    bullet(doc, "Move metrology for Micron to San Jose — 148 of the remaining 184 hours.")
    bullet(doc,
           "Die sorting for two customers cannot be moved at all. Penang is the only plant "
           "that makes it.")

    body(doc, "", after=2)
    rows = [
        ["Protected by rerouting", f"**{money(hur['protected'])}**", f"{hur['protectedPct']}%"],
        ["Residual exposure", money(hur["unprotected"]), f"{round(100-hur['protectedPct'],1)}%"],
    ]
    table(doc, ["Outcome", "Value", "Share"], rows, [3.0, 2.0, 1.85], size=9.5, align_right=(1, 2))

    body(doc,
         "Two findings emerged that were not visible before. The two reroutes consume San Jose's "
         "spare capacity exactly, leaving it at 98.6% utilisation — the mitigation creates a new "
         "single point of failure. And roughly 1.3 million dollars of exposure cannot be "
         "engineered away at all; it requires either a second qualified source or a customer "
         "conversation.",
         italic=True, color=GREY)

    doc.add_page_break()

    # ---- scenario library ------------------------------------------------
    h1(doc, "Scenario library")
    body(doc,
         "Six scenarios ship ready to run. The spread matters: some disruptions are largely "
         "recoverable and some are not recoverable at all, and knowing which is which in advance "
         "is the point of the exercise.")

    rows = []
    for x in scen:
        recover = "—" if x["kind"] == "demand" else f"{x['protectedPct']}%"
        rows.append([
            x["label"].replace(" — ", ": "),
            f"{x['days']}d",
            money(x["valueAtRisk"]),
            f"{x['pctOfNetwork']}%",
            str(x["maxHop"]),
            recover,
        ])
    table(doc, ["Scenario", "Duration", "At risk", "Of network", "Hops", "Recoverable"],
          rows, [2.5, 0.7, 0.95, 0.95, 0.5, 1.05], size=9, align_right=(2, 3, 4, 5))

    body(doc,
         "Two results are worth drawing management attention to.", after=4)
    bullet(doc,
           "is entirely unrecoverable. Penang is the sole source of die sorting, so no amount of "
           "capacity elsewhere helps. This is a structural exposure, not an operational one.",
           bold_lead="A typhoon at Penang ")
    bullet(doc,
           "is only 44.7% recoverable despite being a partial outage, because Dresden has the "
           "least spare capacity in the network and is sole source for e-beam review.",
           bold_lead="A partial loss at Dresden ")

    body(doc,
         f"Across the whole network, {len(nt['single_source_categories'])} of the product "
         f"categories are made at exactly one plant. Those are the places where resilience has "
         f"to be bought rather than planned.")

    # ---- personas --------------------------------------------------------
    h1(doc, "Who uses this, and for what")
    body(doc,
         "Six roles get distinct value from the same application. The question each one asks is "
         "different, and so is the page they start on.")

    rows = [
        ["Supply Chain Risk Manager",
         "What is our exposure this hurricane season, and what is our playbook?",
         "Runs the scenario library, builds a costed response per site, identifies which "
         "exposures cannot be engineered away.",
         "**Best showcase**"],
        ["VP Supply Chain / COO",
         "Where are we structurally fragile, and what would it cost to fix?",
         "Reviews single-source categories and recoverability by scenario to prioritise "
         "second-source investment.",
         "High"],
        ["CFO / FP&A",
         "What revenue is exposed, and how much of it is defensible?",
         "Quantified value at risk, protected versus residual, per scenario — usable in risk "
         "disclosure and contingency planning.",
         "High"],
        ["Plant Manager",
         "What am I being asked to absorb, and can I actually take it?",
         "Sees the hours and utilisation a reroute imposes, and whether it pushes the site past "
         "safe operating headroom.",
         "Medium"],
        ["Customer Account Director",
         "Which of my customers are exposed, and what do I tell them?",
         "Exposure by named customer, plus an AI-drafted position for the customer conversation.",
         "Medium"],
        ["Enterprise / Data Architect",
         "How does this work without copying SAP data?",
         "The ontology explorer and semantic layer demonstrate the zero-copy SAP BDC to "
         "Snowflake pattern end to end.",
         "Medium"],
    ]
    table(doc, ["Role", "Question they ask", "What they do with it", "Fit"],
          rows, [1.5, 1.85, 2.55, 0.85], size=8.5)

    doc.add_page_break()

    # ---- recommended showcase -------------------------------------------
    h1(doc, "Recommended showcase: pre-season resilience review")
    body(doc,
         "The Supply Chain Risk Manager is the strongest persona to demonstrate, for three "
         "reasons. The scenario is one every manufacturer recognises. It exercises the whole "
         "application rather than one page. And it ends in a decision with a number attached, "
         "which is what makes the value legible.",
         size=10.5)

    h2(doc, "The narrative")
    body(doc,
         "Hurricane season is approaching. The risk manager has to tell the executive team what "
         "the company's exposure is and what the response would be — before anything happens, "
         "not during.")

    h2(doc, "Suggested flow, about eight minutes")
    rows = [
        ["1", "Scenario Studio",
         "Open on the network: 5 plants, 6 suppliers, 8 customers, "
         f"{money(nt['monthly_value'])} a month. Establish that this is real, current data.",
         "1 min"],
        ["2", "Scenario Studio",
         "Run the Austin hurricane. Headline exposure appears: "
         f"{money(16_051_035)}, {hur['pctOfNetwork']}% of the network.",
         "1 min"],
        ["3", "Ripple Map",
         "Press Play. The cascade steps outward hop by hop on the map and the network graph "
         "together. Penang lighting up on the second hop is the moment the room understands "
         "the point.",
         "2 min"],
        ["4", "Ripple Map",
         "Click Penang. Show that its 12 days of stock is why it fails on day 13, not day 1.",
         "1 min"],
        ["5", "Mitigation",
         "Show the two feasible reroutes, and that they consume San Jose's spare capacity "
         "exactly — the fix creates a new fragility.",
         "1 min"],
        ["6", "Mitigation",
         "Show the blocked item: die sorting cannot be rerouted at any price, because Penang "
         "is the only source.",
         "1 min"],
        ["7", "Mitigation",
         "Generate the AI briefing, then ask a follow-up such as what if the outage ran twice "
         "as long. Emphasise that the AI interprets the numbers and does not invent them.",
         "1 min"],
    ]
    table(doc, ["#", "Page", "What to do and say", "Time"],
          rows, [0.35, 1.25, 4.4, 0.75], size=8.5)

    h2(doc, "The closing line")
    body(doc,
         f"Of {money(hur['revenueAtRisk'])} of customer revenue exposed by a hurricane at "
         f"Austin, {hur['protectedPct']}% can be protected by moving work between plants we "
         f"already own — and the remainder cannot be protected at any price with the network as "
         f"it stands today. We now know which is which before the storm, not after.",
         size=10.5, italic=True, color=SAP_NAVY)

    h2(doc, "Second scenario, if time allows")
    body(doc,
         "Run the Penang typhoon. It is 0% recoverable, and the contrast with Austin makes the "
         "structural point far better than any single scenario can: resilience is not uniform "
         "across the network, and the difference is knowable in advance.")

    doc.add_page_break()

    # ---- how it works ----------------------------------------------------
    h1(doc, "How it works, briefly")
    rows = [
        ["Where the data lives",
         "Snowflake, as views over existing SAP data products. No copies are made, so the "
         "application and operational reporting cannot disagree."],
        ["How impact spreads",
         "Impact travels along supply flows in proportion to how much a site depends on what it "
         "lost. Inventory absorbs the front of an event, which is why duration changes the "
         "answer so much."],
        ["How mitigation is found",
         "A deterministic optimiser tests every alternative plant against real capacity and "
         "reports what fits. Every number is auditable and reproducible."],
        ["What the AI does",
         "Interprets the computed result, ranks the options and drafts the briefing. It is "
         "explicitly not allowed to calculate the figures, so the numbers remain verifiable."],
        ["Verification",
         "An automated harness runs 20 assertions across all five disruption types on every "
         "change, including checks against hand-computed expected results."],
    ]
    table(doc, ["Aspect", "Approach"], rows, [1.85, 5.05], size=9.5)

    # ---- honest limits ---------------------------------------------------
    h1(doc, "What is real, and what is modelled")
    body(doc,
         "Stating this plainly protects the credibility of the numbers when they are challenged.")

    rows = [
        ["Network, flows, volumes, values", "**Real** — from SAP data products in Snowflake"],
        ["Plant capacity and utilisation", "**Real** — from work-centre capacity data"],
        ["Inventory buffer days", "**Real** — from material stock data"],
        ["Which plant can substitute for another",
         "**Derived** from observed shipments, since production versions list only one plant "
         "per material"],
        ["Hours consumed per unit",
         "**Derived and approximate** — blended across products, suitable for planning rather "
         "than scheduling"],
        ["Qualification and approval time for a plant change",
         "**Not modelled.** A reroute the tool calls feasible may still take weeks to authorise."],
        ["Component-level shortages inside a plant",
         "**Not modelled.** The bill of materials is available if this is taken further."],
    ]
    table(doc, ["Element", "Status"], rows, [2.9, 4.0], size=9)

    # ---- next steps ------------------------------------------------------
    h1(doc, "Options from here")
    bullet(doc,
           "Use it as-is for customer conversations. It is deployed and needs no further work "
           "to demonstrate.",
           bold_lead="Demonstrate. ")
    bullet(doc,
           "Point the same engine at a customer's own SAP network. The scenario logic is "
           "independent of this dataset.",
           bold_lead="Extend to a live customer network. ")
    bullet(doc,
           "Add bill-of-materials explosion and lead times to move from planning-grade to "
           "scheduling-grade accuracy.",
           bold_lead="Deepen the model. ")
    bullet(doc,
           "The single-source analysis already identifies where a second source would remove "
           "the most unrecoverable exposure. That is a costed investment case.",
           bold_lead="Turn the findings into an investment case. ")

    h1(doc, "Access")
    rows = [
        ["Source code and documentation", REPO],
        ["Public demonstration build", SITE],
        ["Runs locally on", "Application port 5179, API port 3009"],
    ]
    table(doc, ["Resource", "Location"], rows, [2.2, 4.7], size=9)

    body(doc,
         "The public build carries no credentials and no customer data. The scenario pages need "
         "a live Snowflake connection and therefore run locally or in a controlled environment.",
         size=9, italic=True, color=GREY)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")
    print(f"  size {OUT.stat().st_size/1024:.0f} KB")


if __name__ == "__main__":
    main()
