#!/usr/bin/env python3
"""Generate per-persona demo scripts as a Word document.

Six self-contained scripts, one per persona, each a page or two so a single
script can be handed to whoever is presenting to that audience. They are
deliberately NOT one template with the role name swapped: each persona opens on a
different page, drives different figures and closes on a different point, because
that is the only way the demo lands as being about their job.

Figures come from /tmp/persona_facts.json, queried from the running application.

Run:  python3 tools/build_demo_scripts.py
"""
import datetime
import json
import pathlib
import sys

from docx import Document
from docx.shared import Pt

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent))
from docx_kit import (  # noqa: E402
    AMBER, GREY, RED, SAP_NAVY, SNOW_BLUE, BLUE_HEX, LIGHT_HEX,
    body, bullet, callout, h1, h2, money, setup_page, table,
)

OUT = pathlib.Path.home() / "Documents" / "SAP" / "Supply_Chain_Ontology_Demo_Scripts.docx"
APP = "http://localhost:5179"


def persona_header(doc, n, role, minutes, audience, question):
    h1(doc, f"Script {n} · {role}", size=17, before=0, after=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"{minutes} minutes   ·   Audience: {audience}")
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY
    callout(doc, "The question they walk in with:", question, fill="EEF4F8", size=10)


def beats(doc, rows):
    """ACTION / SAY table — the presenter reads down the right-hand column."""
    table(doc, ["#", "Page", "Do this", "Say this"], rows,
          [0.3, 1.05, 1.95, 3.6], size=8.5, zebra=True)


def qa(doc, pairs):
    h2(doc, "Questions to expect")
    for q, a in pairs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"“{q}”")
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = SAP_NAVY
        body(doc, a, size=9.5, after=6)


def main():
    try:
        F = json.loads(pathlib.Path("/tmp/persona_facts.json").read_text())
    except FileNotFoundError as e:
        sys.exit(f"missing {e}. Query the running app first.")

    hur, pen = F["hurricane"], F["penang"]
    nt = F["totals"]
    sole = [s for s in F["substitution"] if not s["alt"]]
    sole_value = sum(s["value"] for s in sole)
    sole_pct = round(100 * sole_value / nt["monthly_value"], 1)
    sj = next((c for c in hur["capacityAfter"] if "San Jose" in c["plantName"]), None)
    by_cust = F["hurricane_by_customer"]

    doc = Document()
    setup_page(doc)

    # ---- cover ------------------------------------------------------------
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Supply Chain Ontology")
    r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = SAP_NAVY
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Demo scripts by persona")
    r.font.size = Pt(14); r.font.color.rgb = SNOW_BLUE
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(f"{datetime.date.today():%d %B %Y}   ·   Application: {APP}")
    r.font.size = Pt(9.5); r.font.color.rgb = GREY

    body(doc,
         "Six scripts, each written for one audience. They are not variants of a single "
         "walkthrough — every script opens on a different page, drives different numbers and "
         "closes on a different point, because a demo only lands when it is visibly about the "
         "listener's own job.",
         size=10.5)
    body(doc,
         "Pick the one script that matches the room. Running two back to back works only for "
         "scripts 1 and 2, which are designed to pair.")

    table(doc, ["#", "Persona", "Opens on", "Lands the point that…", "Mins"], [
        ["1", "**Supply Chain Risk Manager**", "Scenario Studio",
         "we can cost a disruption and its response before it happens", "9"],
        ["2", "VP Supply Chain / COO", "Scenario Studio → library",
         f"{sole_pct}% of flow has no alternative source at all", "6"],
        ["3", "CFO / FP&A", "Mitigation",
         "revenue exposure splits into defensible and undefendable", "5"],
        ["4", "Plant Manager", "Mitigation",
         "a reroute is a real load with a real limit", "5"],
        ["5", "Customer Account Director", "Ripple Map",
         "customers are hit who are nowhere near the event", "5"],
        ["6", "Enterprise / Data Architect", "Ontology Graph",
         "none of this required moving SAP data", "7"],
    ], [0.3, 1.9, 1.35, 2.75, 0.55], size=9)

    h2(doc, "Before any demo")
    bullet(doc, f"Application running at **{APP}** — check the Scenario Studio page loads.", size=9.5)
    bullet(doc, "Run the Austin hurricane preset once to warm it, then reset to a clean state.", size=9.5)
    bullet(doc, "For scripts 1, 3 and 5, confirm the AI briefing returns — it needs a live "
                "Snowflake connection.", size=9.5)
    bullet(doc, "Have the second monitor at a readable zoom. The Ripple Map is two panels "
                "side by side and shrinks badly.", size=9.5)

    # =======================================================================
    # 1 — RISK MANAGER
    # =======================================================================
    doc.add_page_break()
    persona_header(
        doc, 1, "Supply Chain Risk Manager", 9, "Risk, resilience and S&OP leads",
        "Hurricane season starts in six weeks. What is our exposure, and what is the plan?")

    body(doc,
         "This is the flagship script. It uses all three pages and ends in a decision with a "
         "number attached. If you only ever learn one script, learn this one.")

    beats(doc, [
        ["1", "Scenario Studio", "Point at the network summary without running anything yet.",
         f"“This is our live network — {nt['plants']} plants, {nt['suppliers']} suppliers, "
         f"{nt['customers']} customers, {money(nt['monthly_value'])} of flow a month. "
         f"It is read from SAP data in Snowflake, not a model someone built in a spreadsheet.”"],
        ["2", "Scenario Studio", "Click the **Hurricane — Austin Fab offline** preset.",
         f"“A hurricane closes Austin for 60 days. {money(hur['risk'])} at risk — "
         f"28.3% of the monthly network. That number took under a second.”"],
        ["3", "Scenario Studio", "Point at the **Exposed** and **Unmitigable** columns.",
         "“The columns are not sorted by size. They are sorted by what each one needs from "
         "me — watch, act, or pick up the phone.”"],
        ["4", "Ripple Map", "Press **Play from the start**. It walks one lane per beat: "
                             "1a, 1b, then 1c.",
         "“1a and 1b are Austin's own customers — no surprise. Watch what 1c does.”"],
        ["4b", "Ripple Map", "Let it land on **1c**, then pause. The map zooms out to the "
                             "whole world on its own.",
         "**Pause here.** “That lane goes to Penang. Penang is in Malaysia — the hurricane "
         "is in Texas. The map just pulled back to show you why that matters.”"],
        ["5", "Ripple Map", "Read the panel above the maps, then click the **Penang** node.",
         "“Austin ships Penang its test fixtures. Penang holds 12 days of stock, so it runs "
         "fine until day 13 and then starves — and takes TSMC and Texas Instruments down "
         "with it. That is the effect no spreadsheet catches.”"],
        ["6", "Mitigation", "Show the two reroutes.",
         f"“Two moves recover {money(hur['protected'])} — {hur['protectedPct']}% of the "
         f"exposure — using plants we already own.”"],
        ["7", "Mitigation", "Point at San Jose's utilisation bar.",
         f"“But look at the cost. San Jose goes to "
         f"{sj['utilizationAfter'] if sj else 98.6}% and has zero spare units left. "
         f"The fix creates a new single point of failure. I would rather know that now than "
         f"discover it in week three.”"],
        ["8", "Mitigation", "Point at the blocked Die Sorting rows.",
         f"“And {money(hur['unprotected'])} cannot be moved at any price. Penang is the only "
         f"plant that makes die sorting. That is not an operations problem, it is an "
         f"investment decision.”"],
        ["9", "Mitigation", "Click **Brief me**, then ask a follow-up.",
         "“The AI reads the numbers the simulation produced — it does not invent them. "
         "Ask it what happens if the outage runs twice as long.”"],
    ])

    h2(doc, "Optional second scenario, 60 seconds")
    body(doc,
         f"Return to Scenario Studio and run **Typhoon — Penang Assembly**. "
         f"{money(pen['risk'])} at risk and **0% recoverable**. "
         f"Say: “Same tool, completely different answer. Austin we can largely absorb. Penang "
         f"we cannot absorb at all. Resilience is not uniform across our network, and now we "
         f"know where it isn't.”", size=9.5)

    callout(doc, "Close on:",
            f"“Of {money(hur['revenue'])} of customer revenue a hurricane at Austin puts at "
            f"risk, {hur['protectedPct']}% is defensible by moving work between plants we "
            f"already own. The rest is not defensible at any price with the network as it "
            f"stands. We now know which is which before the storm, not during it.”",
            fill="EEF4F8", size=10)

    qa(doc, [
        ("Is this real data or a demo dataset?",
         "The network, the volumes, the values, plant capacity and inventory are all real SAP "
         "data products in Snowflake. Two things are derived and the tool labels them: which "
         "plant can substitute for another, and hours consumed per unit."),
        ("How accurate is the 12-day figure?",
         "It is the minimum days of inventory from material stock data. We deliberately use "
         "the minimum, not the average — one component at 12 days stops the line regardless "
         "of another sitting at 400."),
        ("Could we not work this out in a spreadsheet?",
         "The first hop, yes. The second hop is where it breaks down — you have to know that "
         "Austin feeds Penang, how much Penang depends on that flow, and how long Penang's "
         "buffer lasts. That is three joins and a time calculation per affected site."),
    ])

    # =======================================================================
    # 2 — VP / COO
    # =======================================================================
    doc.add_page_break()
    persona_header(
        doc, 2, "VP Supply Chain / COO", 6, "Operations executives and their staff",
        "Where are we structurally fragile, and what would it cost to fix?")

    body(doc,
         "This persona does not want to watch scenarios run. They want the portfolio view: "
         "which exposures are structural, and therefore which ones money can actually fix. "
         "Run fewer scenarios and spend the time on the pattern.")

    beats(doc, [
        ["1", "Scenario Studio", "Run **Hurricane — Austin**, then immediately "
                                 "**Typhoon — Penang**.",
         f"“Two sites, two storms. Austin: {hur['protectedPct']}% of the exposure is "
         f"recoverable. Penang: zero. Same tool, same network.”"],
        ["2", "Scenario Studio", "Stay on the headline numbers. Do not open the map.",
         "“The difference is not severity or duration. It is structure — whether anyone else "
         "can make what that site makes.”"],
        ["3", "Mitigation", "Show the blocked list for Penang.",
         "“Nothing can be rerouted, because Penang is the only plant that makes die sorting. "
         "No amount of spare capacity elsewhere helps.”"],
        ["4", "—", "Present the sole-source table below.",
         f"“Across the network, {len(sole)} of our {len(F['substitution'])} product categories "
         f"are made at exactly one plant. That is {money(sole_value)} a month — {sole_pct}% of "
         f"total flow — with no alternative source.”"],
        ["5", "—", "Let it sit. This is the number they will remember.",
         "“Every one of those lines is a scenario we cannot mitigate our way out of. They are "
         "the shortlist for second-source qualification, and they are ranked by value.”"],
    ])

    h2(doc, "The table to present")
    table(doc, ["Product category", "Plants", "Monthly flow", "Position"],
          [[s["cat"], str(len(s["plants"])), money(s["value"]),
            "**Sole source**" if not s["alt"] else "Has alternative"]
           for s in F["substitution"]],
          [2.5, 0.6, 1.3, 2.5], size=9, align_right=(1, 2), zebra=True)

    callout(doc, "Close on:",
            f"“{sole_pct}% of our monthly flow runs through categories with a single plant "
            f"behind them. Scenario planning cannot fix that — only qualifying a second source "
            f"can. This list is that investment case, ranked by exposure.”",
            fill="EEF4F8", size=10)

    qa(doc, [
        ("Which one would you fix first?",
         f"By exposure, die sorting at {money(5_850_000)} a month. It is also the one that "
         f"already shows up as a second-order failure in the Austin scenario, so it is exposed "
         f"to more than one event."),
        ("What does 'has alternative' really mean — can we actually switch?",
         "It means another plant already ships that category today, so the capability exists. "
         "It does not mean the switch is free: qualification, tooling and customer approval are "
         "not modelled, and the tool says so."),
    ])

    # =======================================================================
    # 3 — CFO
    # =======================================================================
    doc.add_page_break()
    persona_header(
        doc, 3, "CFO / FP&A", 5, "Finance leadership and planning",
        "What revenue is exposed, and how much of it can we actually defend?")

    body(doc,
         "Finance does not need the map. They need the split between defensible and "
         "undefendable exposure, and confidence that the numbers are auditable. Lead with "
         "provenance.")

    beats(doc, [
        ["1", "Mitigation", "Open on a scenario already run. Point at the four tiles.",
         f"“One disruption, four numbers. {money(hur['revenue'])} of customer revenue at risk. "
         f"{money(hur['protected'])} defensible. {money(hur['unprotected'])} not. "
         f"{len(hur['reroutes'])} actions required.”"],
        ["2", "Mitigation", "Point at the reroute table's hours columns.",
         "“Each action is costed in capacity, not guessed. 98 hours of 282 free. The "
         "constraint is real and it is in the data.”"],
        ["3", "Mitigation", "Point at the blocked rows and their reasons.",
         "“This is the part that matters for provisioning. It is not unquantified risk — it is "
         "a specific figure against a specific cause.”"],
        ["4", "Scenario Studio", "Show the scenario library totals.",
         "“Run across our six standing scenarios, exposure ranges from "
         f"{money(3_734_988)} to {money(hur['risk'])}. That is a distribution you can plan "
         f"a contingency against rather than a single worst case.”"],
        ["5", "Mitigation", "Scroll to the limits panel at the foot of the page.",
         "“And the tool states its own assumptions. Anything derived is labelled derived. "
         "I would rather bring you a number with its caveats attached.”"],
    ])

    h2(doc, "The split to present")
    table(doc, ["Scenario", "Revenue at risk", "Defensible", "Undefendable"], [
        ["Hurricane — Austin, 60d", money(hur["revenue"]),
         f"{money(hur['protected'])} ({hur['protectedPct']}%)", money(hur["unprotected"])],
        ["Typhoon — Penang, 30d", money(pen["revenue"]), "$0 (0%)", money(pen["revenue"])],
    ], [2.5, 1.5, 1.6, 1.3], size=9, align_right=(1, 2, 3))

    callout(doc, "Close on:",
            "“Exposure is now two numbers instead of one: what we can defend by moving work, "
            "and what we cannot defend at all. The second number is the one that belongs in a "
            "contingency conversation, and it is defensible line by line.”",
            fill="EEF4F8", size=10)

    qa(doc, [
        ("Can I put this in a risk disclosure?",
         "The methodology is documented and reproducible, and every figure traces to SAP data "
         "in Snowflake. Treat the derived elements as planning-grade — the tool marks them."),
        ("Why does the same plant give different numbers at different durations?",
         "Because inventory absorbs the front of an event. A 10-day Austin outage barely "
         "reaches Penang; a 60-day one takes it out for 48 days. Duration is the single "
         "biggest driver in the model."),
    ])

    # =======================================================================
    # 4 — PLANT MANAGER
    # =======================================================================
    doc.add_page_break()
    persona_header(
        doc, 4, "Plant Manager", 5, "Site leadership and production planning",
        "If another site goes down, what am I being asked to absorb — and can I?")

    body(doc,
         "This persona is on the receiving end. The demo works because it shows the ask in "
         "their units — hours and utilisation — and because it is honest about when the answer "
         "is no. Present it from San Jose's point of view.")

    beats(doc, [
        ["1", "Scenario Studio", "Run **Hurricane — Austin**. Scroll to the capacity bullets.",
         "“Austin is down. Before anything is decided, here is where every plant stands "
         "today.”"],
        ["2", "Scenario Studio", "Point at the band and marker on the bars.",
         "“Grey band is the normal operating range, the line is full capacity. Not a dial — "
         "you can read the actual distance to the limit.”"],
        ["3", "Mitigation", "Show the two reroutes going to San Jose.",
         f"“San Jose is asked to take {sj['unitsAdded'] if sj else 5} units a month across two "
         f"product lines — {sj['hrsAdded'] if sj else 245.8} hours.”"],
        ["4", "Mitigation", "Point at the before-and-after utilisation.",
         f"“That moves San Jose from {sj['utilizationBefore'] if sj else 89.1}% to "
         f"{sj['utilizationAfter'] if sj else 98.6}%, with zero spare units left. The tool "
         f"flags it: above 95% there is no recovery room if anything else slips.”"],
        ["5", "Scenario Studio", "Run **Partial loss — Dresden Fab at 60%**.",
         "“Now the opposite answer. Dresden has the least headroom in the network, so when it "
         "is the one in trouble only 44.7% of the exposure can be moved.”"],
        ["6", "Mitigation", "Point at a blocked row citing no spare capacity.",
         "“When the answer is no, it says no and tells you how many units short. That is a "
         "conversation about overtime or displacing lower-value work — not a surprise in "
         "week two.”"],
    ])

    h2(doc, "Current position across the network")
    table(doc, ["Plant", "Utilisation", "Free hours", "Spare units/mo", "Min stock"],
          [[c["plant_name"], f"{c['utilization_pct']}%", str(c["free_hrs"]),
            str(c["spare_units"]),
            next((f"{i['min_days_of_inventory']}d" for i in F["inventory"]
                  if i["plant"] == c["plant"]), "—")]
           for c in sorted(F["capacity"], key=lambda x: -(x["spare_units"] or 0))],
          [1.9, 1.15, 1.05, 1.4, 1.0], size=9, align_right=(1, 2, 3, 4), zebra=True)

    callout(doc, "Close on:",
            "“When a reroute lands on your site you get the ask in hours, the resulting "
            "utilisation, and a warning if it leaves you with no room. And when your site is "
            "genuinely full, the plan says so instead of assuming you will absorb it.”",
            fill="EEF4F8", size=10)

    qa(doc, [
        ("Hours per unit for my plant is not that simple.",
         "Agreed, and the tool labels that figure as derived — used hours divided by units "
         "shipped, blended across products. It is right for 'can this site absorb five more "
         "units', not for building a routing."),
        ("Does it check I can get the components?",
         "No. Inbound component availability at the receiving plant is not re-checked, and "
         "that is listed as a limitation. It is the first thing to verify before committing."),
    ])

    # =======================================================================
    # 5 — ACCOUNT DIRECTOR
    # =======================================================================
    doc.add_page_break()
    persona_header(
        doc, 5, "Customer Account Director", 5, "Sales and account management",
        "Which of my customers are affected, and what do I tell them this week?")

    body(doc,
         "Lead with the customer list, not the network. The hook for this persona is that two "
         "of the four affected customers have no connection to Texas at all.")

    beats(doc, [
        ["1", "Ripple Map", "Run the Austin hurricane, then scroll to the flows table.",
         "“Austin is closed by a hurricane. Four of my customers are affected.”"],
        ["2", "Ripple Map", "Step to **1a** and **1b** as you name them.",
         f"“{by_cust[0]['customer']} and {by_cust[1]['customer']} are hop one — they buy from "
         f"Austin directly, no surprise there.”"],
        ["3", "Ripple Map", "Step to **2a** and **2b**. Slow down here.",
         f"“{by_cust[2]['customer']} and {by_cust[3]['customer']} are hop two. Neither buys "
         f"anything from Austin. They buy from Penang — and Penang depends on Austin. "
         f"Without this I would not have called them.”"],
        ["4", "Ripple Map", "Click Penang and read the buffer line in the step panel.",
         "“And I know the timing. Penang holds 12 days of stock, so those two accounts are "
         "fine for a fortnight. That changes whether I call today or next week.”"],
        ["5", "Mitigation", "Show which customers the reroutes protect.",
         f"“{by_cust[0]['customer']} and {by_cust[1]['customer']} I can protect by moving "
         f"production to San Jose. So those are reassurance calls.”"],
        ["6", "Mitigation", "Point at the blocked die sorting rows.",
         f"“{by_cust[2]['customer']} and {by_cust[3]['customer']} I cannot protect — nobody "
         f"else makes die sorting. Those are honest-conversation calls, and I would rather "
         f"make them early.”"],
        ["7", "Mitigation", "Ask the AI: “What should I tell the affected customers?”",
         "“And it will draft the position for each one from the same numbers.”"],
    ])

    h2(doc, "Your call list, in priority order")
    table(doc, ["Customer", "Exposure", "Hop", "Product", "The call"],
          [[c["customer"], money(c["value"]), str(c["hop"]), ", ".join(c["cats"]),
            "Reassure — protected by reroute" if c["hop"] == 1
            else "**Early warning — cannot be protected**"]
           for c in by_cust],
          [1.45, 0.85, 0.4, 1.35, 2.4], size=9, align_right=(1, 2), zebra=True)

    callout(doc, "Close on:",
            "“Two of these four customers are nowhere near the hurricane. I would not have "
            "known to call them, and I would certainly not have known which ones I could "
            "reassure and which ones I could not.”",
            fill="EEF4F8", size=10)

    qa(doc, [
        ("How confident are the dollar figures per customer?",
         "They are the monthly flow value for that customer, prorated over the days actually "
         "exposed. The flow values come from SAP data, so the split between customers is as "
         "good as our own order data."),
        ("Can I get this as a report to take to the account team?",
         "The figures come from an API, so yes — the same numbers can be exported. Today the "
         "screen is the deliverable."),
    ])

    # =======================================================================
    # 6 — ARCHITECT
    # =======================================================================
    doc.add_page_break()
    persona_header(
        doc, 6, "Enterprise / Data Architect", 7, "Data platform, integration and EA",
        "How does this work, and what did it cost us in data movement and new infrastructure?")

    body(doc,
         "This is the only script that starts on the ontology pages. The point to land is that "
         "the scenario capability is a thin layer over data that was already in place — no "
         "pipeline, no copies, no new platform.")

    beats(doc, [
        ["1", "Ontology Graph", "Open the graph. Filter to one process.",
         "“Start with what SAP gives us. Business Data Cloud publishes the supply chain as "
         "governed data products, each one a set of CDS entities with declared "
         "relationships.”"],
        ["2", "Ontology Graph", "Expand a data product to show its entities.",
         "“36 supply-chain data products, 338 entities, 281 declared associations. The "
         "relationships are in SAP's own metadata — we did not reverse-engineer them from "
         "column names.”"],
        ["3", "Graph Traversal", "Show the amber topology panel.",
         "“And we measured the shape rather than assuming it. No association crosses a data "
         "product boundary; cross-product linkage runs entirely through one canonical object, "
         "Plant. That is a finding, and the page says so.”"],
        ["4", "Ask the Ontology", "Ask a question and expand the generated SQL.",
         "“Natural language over a governed semantic view, and it shows you the SQL. The "
         "answer is auditable, not asserted.”"],
        ["5", "Scenario Studio", "Switch to the scenario section.",
         f"“Now the scenario layer. {nt['nodes']} nodes, {nt['flows']} flows, "
         f"{money(nt['monthly_value'])} a month — all views over existing Snowflake tables. "
         f"No copies, so this and operational reporting cannot disagree.”"],
        ["6", "Ripple Map", "Run a scenario and note the response time.",
         "“Propagation runs in the application, not the warehouse. Nineteen nodes and 27 "
         "flows simulate in microseconds, so the UI recomputes on every slider move. A "
         "warehouse round-trip per keystroke would buy nothing.”"],
        ["7", "Mitigation", "Click Brief me.",
         "“The narrative uses AI_COMPLETE, deliberately not Cortex Analyst. The ripple exists "
         "only in memory — Analyst would query the undisrupted network and confidently answer "
         "the wrong question.”"],
        ["8", "—", "Summarise the build.",
         "“Views, one semantic view, a static JSON export and an application. No ETL, no new "
         "platform, and a verification harness that runs 20 assertions across five disruption "
         "types on every change.”"],
    ])

    h2(doc, "What was actually built")
    table(doc, ["Layer", "What exists", "Data movement"], [
        ["SAP BDC → Snowflake", "Existing data products", "**None** — already in place"],
        ["Ontology", "6 views + 1 semantic view over shared CORE tables", "**None** — views only"],
        ["Scenario network", "5 views over SAP_SUPPLY_CHAIN", "**None** — views only"],
        ["Served artifact", "One 23 KB JSON export, regenerated on demand", "Export only"],
        ["Application", "Node API + React client", "Reads the JSON"],
        ["AI", "AI_COMPLETE for narrative; Analyst for the ontology", "Prompt only"],
    ], [1.75, 3.35, 1.8], size=9, zebra=True)

    callout(doc, "Close on:",
            "“The expensive part — getting governed SAP data into Snowflake with its semantics "
            "intact — was already done. Scenario modelling turned out to be a thin layer on "
            "top of it, which is the whole argument for doing the foundation properly.”",
            fill="EEF4F8", size=10)

    qa(doc, [
        ("Why keep the simulation out of Snowflake?",
         "Interactivity. At 19 nodes the graph fits in memory and a scenario is microseconds. "
         "If the network were thousands of nodes the answer would flip to graph processing in "
         "the warehouse."),
        ("What happens when the SAP catalog changes?",
         "Three scripted commands re-slice the ontology, re-assert Snowflake/JSON parity and "
         "re-snapshot the public build. The parity check fails the run if the two "
         "implementations of the scope rule have drifted."),
        ("Could this point at a customer's network instead?",
         "Yes. The scenario logic reads a network contract, not this dataset. The work would be "
         "mapping their tables to the five views."),
    ])

    doc.save(OUT)
    print(f"wrote {OUT}")
    print(f"  size {OUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
